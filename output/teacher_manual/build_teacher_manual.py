from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
OUT = ROOT / "AI_로봇_코치_교사용_운영_설명서.docx"

NAVY = "102A56"
BLUE = "4D58E8"
INDIGO = "3047C8"
LIGHT = "EEF2FF"
PALE = "F7F9FD"
LINE = "D9E1EF"
TEXT = "17243A"
MUTED = "66758F"
WHITE = "FFFFFF"
GREEN = "149A67"
ORANGE = "D97917"
RED = "C64646"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in edges:
            continue
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edges[edge].items():
            element.set(qn("w:" + key), str(value))


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn("w:" + tag))
        if node is None:
            node = OxmlElement("w:" + tag)
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    tr_pr.append(el)


def set_run_font(run, name="맑은 고딕", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text, *, size=9, bold=False, color=TEXT, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    return p


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run()
    set_run_font(r, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    r._r.append(fld_char1)
    r._r.append(instr)
    r._r.append(fld_char2)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn("w:" + edge))
        if el is None:
            el = OxmlElement("w:" + edge)
            borders.append(el)
        el.set(qn("w:val"), "nil")


def set_table_widths(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


doc = Document()
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(18)
section.bottom_margin = Mm(17)
section.left_margin = Mm(18)
section.right_margin = Mm(18)
section.different_first_page_header_footer = True

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "맑은 고딕"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
normal.font.size = Pt(9.5)
normal.font.color.rgb = RGBColor.from_string(TEXT)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.22

for style_name, size, color, before, after in [
    ("Title", 26, NAVY, 0, 8),
    ("Subtitle", 12, BLUE, 0, 6),
    ("Heading 1", 19, NAVY, 0, 9),
    ("Heading 2", 13, INDIGO, 8, 5),
    ("Heading 3", 10.5, NAVY, 5, 3),
]:
    st = styles[style_name]
    st.font.name = "맑은 고딕"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

if "Caption KR" not in styles:
    cap = styles.add_style("Caption KR", WD_STYLE_TYPE.PARAGRAPH)
else:
    cap = styles["Caption KR"]
cap.font.name = "맑은 고딕"
cap._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
cap.font.size = Pt(8)
cap.font.color.rgb = RGBColor.from_string(MUTED)
cap.paragraph_format.space_before = Pt(3)
cap.paragraph_format.space_after = Pt(5)


def header_footer(sec):
    sec.header_distance = Mm(7)
    sec.footer_distance = Mm(7)
    h = sec.header
    p = h.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("생각을 키우는 AI 코치 활용 로봇 수업 자료  |  교사용 운영 설명서")
    set_run_font(r, size=7.5, bold=True, color=MUTED)
    f = sec.footer
    t = f.add_table(rows=1, cols=2, width=Mm(174))
    remove_table_borders(t)
    set_cell_text(t.cell(0, 0), "운천고등학교 · 유동규 · 2026학년도", size=7.5, color=MUTED)
    set_page_number(t.cell(0, 1).paragraphs[0])
    set_cell_border(t.cell(0, 0), top={"val": "single", "sz": "6", "color": LINE})
    set_cell_border(t.cell(0, 1), top={"val": "single", "sz": "6", "color": LINE})


header_footer(section)


def title(text, kicker=None, intro=None):
    if kicker:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(kicker.upper())
        set_run_font(r, size=8, bold=True, color=BLUE)
    doc.add_paragraph(text, style="Heading 1")
    if intro:
        p = doc.add_paragraph(intro)
        p.paragraph_format.space_after = Pt(8)
        for r in p.runs:
            set_run_font(r, size=9.5, color=MUTED)


def p(text, *, bold_prefix=None, color=None, size=None, align=None, space_after=5):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = para.add_run(bold_prefix)
        set_run_font(r1, size=size or 9.5, bold=True, color=color or TEXT)
        r2 = para.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=size or 9.5, color=color or TEXT)
    else:
        r = para.add_run(text)
        set_run_font(r, size=size or 9.5, color=color or TEXT)
    return para


def bullets(items, *, numbered=False, level=0, size=9):
    style = "Normal" if numbered else ("List Bullet 2" if level else "List Bullet")
    for idx, item in enumerate(items, start=1):
        para = doc.add_paragraph(style=style)
        if numbered:
            para.paragraph_format.left_indent = Cm(0.45)
            para.paragraph_format.first_line_indent = Cm(-0.45)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = 1.15
        r = para.add_run(f"{idx}.  {item}" if numbered else item)
        set_run_font(r, size=size, color=TEXT)


def callout(label, text, color=BLUE, fill=LIGHT):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_table_widths(t, [2.2, 14.6])
    remove_table_borders(t)
    set_cell_shading(t.cell(0, 0), color)
    set_cell_shading(t.cell(0, 1), fill)
    set_cell_text(t.cell(0, 0), label, size=8.5, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(t.cell(0, 1), text, size=8.5, color=TEXT)
    for cell in t.rows[0].cells:
        set_cell_border(cell, top={"val": "single", "sz": "4", "color": color},
                        bottom={"val": "single", "sz": "4", "color": color})
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def cards(items, columns=3):
    rows = (len(items) + columns - 1) // columns
    t = doc.add_table(rows=rows, cols=columns)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    width = 16.8 / columns
    set_table_widths(t, [width] * columns)
    remove_table_borders(t)
    for idx in range(rows * columns):
        cell = t.cell(idx // columns, idx % columns)
        set_cell_margins(cell, top=160, start=160, bottom=160, end=160)
        set_cell_border(cell, top={"val": "single", "sz": "5", "color": LINE},
                        left={"val": "single", "sz": "5", "color": LINE},
                        right={"val": "single", "sz": "5", "color": LINE},
                        bottom={"val": "single", "sz": "5", "color": LINE})
        if idx >= len(items):
            set_cell_shading(cell, WHITE)
            continue
        head, body = items[idx]
        set_cell_shading(cell, PALE)
        cell.text = ""
        hp = cell.paragraphs[0]
        hp.paragraph_format.space_after = Pt(5)
        rr = hp.add_run(head)
        set_run_font(rr, size=10, bold=True, color=INDIGO)
        bp = cell.add_paragraph()
        bp.paragraph_format.space_after = Pt(0)
        br = bp.add_run(body)
        set_run_font(br, size=8.3, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def simple_table(headers, rows, widths=None, font_size=8.2):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    if widths:
        set_table_widths(t, widths)
    for i, h in enumerate(headers):
        set_cell_shading(t.cell(0, i), NAVY)
        set_cell_text(t.cell(0, i), h, size=8.2, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(t.rows[0])
    for row_data in rows:
        cells = t.add_row().cells
        cant_split(t.rows[-1])
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], str(value), size=font_size)
            set_cell_border(cells[i], bottom={"val": "single", "sz": "4", "color": LINE})
            if len(t.rows) % 2 == 1:
                set_cell_shading(cells[i], PALE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def process_flow(items):
    t = doc.add_table(rows=1, cols=len(items) * 2 - 1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    widths = []
    arrow_width = 0.55
    box_width = (16.8 - arrow_width * (len(items) - 1)) / len(items)
    for i in range(len(items) * 2 - 1):
        widths.append(box_width if i % 2 == 0 else arrow_width)
    set_table_widths(t, widths)
    remove_table_borders(t)
    for i, item in enumerate(items):
        c = t.cell(0, i * 2)
        set_cell_shading(c, LIGHT if i % 2 == 0 else PALE)
        set_cell_border(c, top={"val": "single", "sz": "5", "color": BLUE},
                        left={"val": "single", "sz": "5", "color": BLUE},
                        right={"val": "single", "sz": "5", "color": BLUE},
                        bottom={"val": "single", "sz": "5", "color": BLUE})
        set_cell_text(c, item, size=8, bold=True, color=INDIGO, align=WD_ALIGN_PARAGRAPH.CENTER)
        if i < len(items) - 1:
            set_cell_text(t.cell(0, i * 2 + 1), "→", size=11, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def screenshot(name, caption):
    path = ASSETS / name
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(path), width=Cm(16.7))
    cp = doc.add_paragraph(caption, style="Caption KR")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def page_break():
    doc.add_page_break()


# 1. Cover
cover_band = doc.add_table(rows=1, cols=1)
cover_band.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_band.autofit = False
set_table_widths(cover_band, [17.4])
set_cell_shading(cover_band.cell(0, 0), NAVY)
set_cell_text(cover_band.cell(0, 0), "제73회 교육자료전 · 2026학년도", size=10, bold=True, color=WHITE)
doc.add_paragraph().paragraph_format.space_after = Pt(30)
pp = doc.add_paragraph()
pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
rr = pp.add_run("생각을 키우는")
set_run_font(rr, size=19, bold=True, color=BLUE)
pp.paragraph_format.space_after = Pt(2)
pp = doc.add_paragraph()
rr = pp.add_run("AI 코치 활용\n로봇 수업 자료")
set_run_font(rr, size=29, bold=True, color=NAVY)
pp.paragraph_format.line_spacing = 1.0
pp.paragraph_format.space_after = Pt(12)
pp = doc.add_paragraph()
rr = pp.add_run("교사용 운영 설명서")
set_run_font(rr, size=15, bold=True, color=INDIGO)
pp.paragraph_format.space_after = Pt(34)

cover_cards = doc.add_table(rows=1, cols=3)
cover_cards.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_cards.autofit = False
set_table_widths(cover_cards, [5.4, 5.4, 5.4])
remove_table_borders(cover_cards)
for i, (head, body) in enumerate([
    ("수업의 중심", "정답보다 질문"),
    ("학습의 과정", "실험·기록·성찰"),
    ("교사의 역할", "관찰·피드백·지원"),
]):
    c = cover_cards.cell(0, i)
    set_cell_shading(c, LIGHT)
    set_cell_border(c, top={"val": "single", "sz": "8", "color": BLUE})
    c.text = ""
    a = c.paragraphs[0]
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a.paragraph_format.space_after = Pt(4)
    r = a.add_run(head)
    set_run_font(r, size=8, bold=True, color=MUTED)
    b = c.add_paragraph()
    b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = b.add_run(body)
    set_run_font(r, size=11, bold=True, color=NAVY)
doc.add_paragraph().paragraph_format.space_after = Pt(35)
meta = doc.add_table(rows=3, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.RIGHT
meta.autofit = False
set_table_widths(meta, [3.0, 6.4])
remove_table_borders(meta)
for i, (a, b) in enumerate([("소속", "운천고등학교"), ("제작", "유동규"), ("운영 연도", "2026학년도")]):
    set_cell_text(meta.cell(i, 0), a, size=9, bold=True, color=MUTED)
    set_cell_text(meta.cell(i, 1), b, size=10, bold=True, color=NAVY)
    set_cell_border(meta.cell(i, 0), bottom={"val": "single", "sz": "4", "color": LINE})
    set_cell_border(meta.cell(i, 1), bottom={"val": "single", "sz": "4", "color": LINE})
page_break()

# 2. At a glance
title("한눈에 보는 자료", "AT A GLANCE", "학생의 실물 로봇 활동과 AI 질문 코칭, 과정 기록, 교사의 분석·피드백을 하나의 수업 흐름으로 연결한 웹 기반 수업 자료입니다.")
cards([
    ("대상", "중·고등학교 로봇·정보·기술 수업"),
    ("형태", "웹앱 + 실물 키트 + 가상 시뮬레이션"),
    ("핵심 활동", "아두이노·로봇팔·메카넘 실습"),
    ("사고 지원", "질문형 AI 코치와 AI 컴파일"),
    ("과정 기록", "엔지니어링 노트와 질문 이력"),
    ("교사 지원", "대시보드·PDF·개별/전체 분석"),
], 3)
doc.add_paragraph("핵심 수업 흐름", style="Heading 2")
process_flow(["미션", "구성", "실행", "AI 질문", "기록", "피드백"])
doc.add_paragraph("이 자료가 해결하려는 수업 문제", style="Heading 2")
simple_table(
    ["수업에서 자주 보이는 문제", "자료의 대응 방식"],
    [
        ("예제 코드를 그대로 복사한다.", "참고 조건과 미션 조건을 다르게 제시해 비교·수정하게 한다."),
        ("오류가 나면 곧바로 정답을 묻는다.", "AI가 답 대신 확인 순서와 비교 질문을 제시한다."),
        ("결과만 남고 사고 과정은 사라진다.", "엔지니어링 노트에 문제·원인·해결·성찰을 구조화한다."),
        ("교사가 모든 학생의 과정을 동시에 보기 어렵다.", "질문·노트·진행률·분석을 교사용 화면에 모은다."),
    ],
    [7.5, 9.3],
)
callout("핵심", "AI는 학생의 정답 생성기가 아니라, 관찰과 비교를 촉진하는 질문형 코치로 사용합니다.", GREEN, "EAF8F2")
page_break()

# 3. TOC
title("차례", "CONTENTS", "수업 전에 3장까지 확인하고, 실제 운영 중에는 필요한 장을 빠르게 찾아 활용할 수 있도록 구성했습니다.")
toc_rows = [
    ("1", "자료 소개", "4"),
    ("2", "시스템 이해", "5"),
    ("3", "수업 시작 준비", "6–7"),
    ("4", "학생 실습 기능", "8–12"),
    ("5", "AI 코치 활용법", "13"),
    ("6", "교사용 대시보드", "14–15"),
    ("7", "피드백과 결과물", "16"),
    ("8", "AI 학습 분석", "17–18"),
    ("9", "관리자 기능", "19"),
    ("10", "실제 수업 운영안", "20–21"),
    ("11", "평가 및 관찰", "22"),
    ("12", "안전·개인정보·문제 해결", "23–24"),
    ("부록", "체크리스트·FAQ·빠른 참조", "25–26"),
]
simple_table(["구분", "내용", "쪽"], toc_rows, [2.0, 12.5, 2.3], 9)
callout("안내", "화면 이미지는 사용법 설명을 위해 제작한 비식별 예시입니다. 실제 화면은 계정 권한과 저장 데이터에 따라 달라질 수 있습니다.", BLUE, LIGHT)
page_break()

# 4. Intro
title("1. 자료 소개", "01 · PURPOSE", "실물 로봇 수업의 ‘만들기’ 경험을 문제 발견, 원인 분석, 개선, 성찰의 학습 경험으로 확장합니다.")
doc.add_paragraph("개발 목적", style="Heading 2")
p("학생은 회로와 기구를 구성하고 코드를 실행하는 과정에서 필연적으로 오류와 예상 밖의 결과를 만납니다. 이 자료는 그 순간을 실패가 아니라 탐구의 출발점으로 바꾸기 위해 설계되었습니다.")
cards([
    ("관찰", "무엇이 어떻게 동작했는지 사실대로 확인"),
    ("비교", "예상과 실제, 참고 예제와 미션 조건을 대조"),
    ("가설", "원인 후보를 한 가지씩 말로 표현"),
    ("검증", "값 하나를 바꾸고 결과를 다시 확인"),
    ("기록", "바꾼 값과 결과를 엔지니어링 노트에 남김"),
    ("성찰", "다음 실험의 기준과 개선 계획을 세움"),
], 3)
doc.add_paragraph("교육적 특징", style="Heading 2")
bullets([
    "실물 키트와 가상 시뮬레이션을 번갈아 활용해 추상적인 원리를 눈에 보이게 합니다.",
    "정답 코드 대신 조건이 다른 참고 코드를 제시해 학생이 차이를 찾아 수정하게 합니다.",
    "AI 질문, 노트, 진행률을 기록해 결과뿐 아니라 문제 해결 과정을 평가할 수 있습니다.",
    "교사는 개별 학생과 전체 학급의 어려움 신호를 빠르게 파악해 다음 피드백을 설계할 수 있습니다.",
])
callout("교사 관점", "성공 여부보다 ‘무엇을 관찰했고, 어떤 값을 바꾸었으며, 결과를 어떻게 설명하는가’를 수업의 주요 증거로 봅니다.", INDIGO, LIGHT)
page_break()

# 5. System
title("2. 시스템 이해", "02 · SYSTEM", "학생용 학습 공간, 교사용 관찰·피드백 공간, 관리자 설정이 역할별로 분리되어 있습니다.")
doc.add_paragraph("역할별 화면", style="Heading 2")
simple_table(
    ["역할", "주요 기능", "데이터 범위"],
    [
        ("학생", "실습, AI 코치, 엔지니어링 노트, 진행 저장", "본인의 학습 기록"),
        ("교사", "학생 관리, 질문/노트 확인, PDF, AI 분석", "같은 학교 학생"),
        ("관리자", "기능 켜기/끄기, GPT 모델, 역할 관리", "시스템 설정과 전체 회원"),
    ],
    [2.5, 8.8, 5.5],
)
doc.add_paragraph("서비스 흐름", style="Heading 2")
process_flow(["브라우저", "Firebase 인증", "Firestore 기록", "Functions 중계", "OpenAI API"])
p("학생과 교사는 Google 계정으로 로그인합니다. 학습 기록은 Firebase에 저장되고, AI 요청은 브라우저에서 API 키를 직접 사용하지 않고 서버 기능을 거쳐 처리됩니다.")
doc.add_paragraph("주요 기술 구성", style="Heading 2")
cards([
    ("웹 인터페이스", "HTML·CSS·JavaScript 모듈"),
    ("시뮬레이션", "Canvas API·Three.js"),
    ("인증·저장", "Firebase Authentication·Firestore"),
    ("AI 중계", "Firebase Functions·Secret Manager"),
    ("배포", "Firebase Hosting·GitHub Actions"),
    ("접속 주소", "robot-ai-class.web.app"),
], 3)
callout("보안", "OpenAI API 키는 학생 브라우저나 문서에 포함하지 않습니다. 관리자는 키가 아닌 기능 상태와 사용할 모델을 관리합니다.", GREEN, "EAF8F2")
page_break()

# 6. Prep
title("3. 수업 시작 준비", "03 · BEFORE CLASS", "처음 사용하는 차시에는 로그인·권한·기기·키트를 먼저 확인하고, 수업 중 확인할 증거를 미리 정합니다.")
doc.add_paragraph("교사 사전 체크리스트", style="Heading 2")
simple_table(
    ["확인", "점검 항목", "완료 기준"],
    [
        ("□", "교사 계정 로그인", "교사용 대시보드 메뉴가 보임"),
        ("□", "학생 계정/학교명/학번", "학생 관리 목록에서 검색됨"),
        ("□", "AI 코치 사용 상태", "관리자 메뉴에서 켜짐 확인"),
        ("□", "브라우저", "Chrome 또는 Edge 최신 버전"),
        ("□", "인터넷·방화벽", "로그인, 저장, AI 응답 정상"),
        ("□", "키트·케이블·전원", "팀별 부품 수량과 이상 유무 확인"),
        ("□", "수업 목표", "완성 결과와 과정 증거를 각각 한 문장으로 설정"),
    ],
    [1.2, 7.2, 8.4],
)
doc.add_paragraph("권장 환경", style="Heading 2")
cards([
    ("교사용", "노트북/PC, 프로젝터, 안정적 네트워크"),
    ("학생용", "팀당 PC 1대, 실물 키트 1세트"),
    ("브라우저", "Chrome·Edge 권장, 인앱 브라우저 제외"),
], 3)
callout("주의", "카카오톡·네이버·인스타그램 안에서 열린 브라우저는 Google 로그인이 제한될 수 있습니다. 주소를 복사해 Chrome 또는 Edge에서 여세요.", ORANGE, "FFF5E8")
p("접속 주소: https://robot-ai-class.web.app", bold_prefix="접속 주소:", color=NAVY)
page_break()

# 7. Login roles
title("3. 수업 시작 준비", "03 · LOGIN & ROLE", "로그인 후 역할과 학교 정보가 올바르게 연결되어야 교사용 데이터가 정확히 모입니다.")
doc.add_paragraph("첫 로그인 절차", style="Heading 2")
bullets([
    "웹앱 접속 후 Google 로그인을 선택합니다.",
    "학교명, 학번, 이름 등 기본 정보를 입력하거나 확인합니다.",
    "학생은 학생 실습실, 교사는 교사용 대시보드가 보이는지 확인합니다.",
    "교사는 학생 관리에서 같은 학교 학생이 검색되는지 점검합니다.",
], numbered=True)
doc.add_paragraph("역할이 잘못 표시될 때", style="Heading 2")
simple_table(
    ["증상", "확인", "조치"],
    [
        ("교사 메뉴가 보이지 않음", "프로필 역할", "관리자에게 교사 역할 지정 요청"),
        ("학생이 교사 목록에 없음", "학교명 철자", "학생·교사 학교명을 동일하게 수정"),
        ("최근 활동이 없음", "로그인 계정", "학생이 같은 계정으로 실습·저장했는지 확인"),
        ("Google 로그인 실패", "접속 브라우저", "외부 Chrome/Edge에서 다시 접속"),
    ],
    [5.1, 4.2, 7.5],
)
callout("권한 원칙", "관리자 메뉴는 지정된 관리자 계정 권한으로 접근합니다. 공용 초기 비밀번호를 배포하거나 문서에 기록하지 않습니다.", RED, "FDEEEE")
doc.add_paragraph("수업 시작 안내 문장", style="Heading 2")
callout("예시", "오늘은 완성 속도보다 ‘어떤 차이를 발견했고 무엇을 바꾸었는지’를 기록합니다. AI에게는 정답을 요구하기보다 다음에 확인할 항목을 질문하세요.", BLUE, LIGHT)
page_break()

# 8. Student dashboard
title("4. 학생 실습 기능", "04 · STUDENT DASHBOARD", "학생은 대시보드에서 학습할 모듈을 선택하고 완료 레슨, AI 질문, 노트 기록을 확인합니다.")
screenshot("screen_overview.png", "그림 1. 학생 대시보드 비식별 예시")
simple_table(
    ["화면 요소", "교사 안내 포인트"],
    [
        ("완료 레슨", "차시 종료 전 완료 처리를 확인하되, 완성 여부와 학습 품질을 동일시하지 않음"),
        ("AI 질문", "질문의 수보다 질문이 관찰·비교·검증으로 발전하는지 확인"),
        ("엔지니어링 노트", "수업마다 최소 1개의 문제 해결 기록을 남기도록 안내"),
        ("최근 활동", "다음 차시에 이어서 할 위치를 빠르게 확인"),
    ],
    [4.0, 12.8],
    8,
)
callout("교사 확인", "차시 시작 3분 동안 모든 학생이 올바른 계정으로 로그인했고 오늘의 모듈에 들어갔는지 확인합니다.", GREEN, "EAF8F2")
page_break()

# 9. Arduino
title("4. 학생 실습 기능", "04 · ARDUINO LAB", "아두이노 실습은 미션 이해, 회로 구성, 코드 수정, 가상 실행, 실제 키트 검증의 순서로 운영합니다.")
screenshot("screen_arduino.png", "그림 2. 아두이노 실습실과 질문형 AI 코치 비식별 예시")
doc.add_paragraph("단원 구성", style="Heading 2")
simple_table(
    ["단원", "핵심 개념", "대표 관찰"],
    [
        ("LED", "디지털 출력·핀 번호", "회로 핀과 코드 핀의 일치"),
        ("서보모터", "각도·PWM", "목표 각도와 실제 회전"),
        ("DC모터", "방향·속도", "극성·드라이버·속도 값"),
        ("초음파센서", "거리 측정", "단위·임계값·반응"),
        ("블루투스", "문자 명령·통신", "입력값과 동작의 연결"),
    ],
    [3.2, 5.5, 8.1],
    7.8,
)
callout("수업 팁", "참고 코드를 그대로 복사하게 하지 말고, ‘참고 조건과 오늘 미션에서 달라진 한 가지’를 먼저 말하게 합니다.", BLUE, LIGHT)
page_break()

# 10. Robot arm
title("4. 학생 실습 기능", "04 · ROBOT ARM", "로봇팔 실습은 관절 각도와 끝점 좌표의 관계를 탐색하고, 경로와 충돌을 고려한 움직임을 설계합니다.")
cards([
    ("2D 로봇팔", "숄더·엘보 각도, 5 cm 좌표계, 끝점 좌표, 도달 가능 범위"),
    ("3D 로봇팔", "베이스·숄더·엘보 조절, X·Y·Z 좌표, 장애물과 경로"),
    ("물건 옮기기", "접근–하강–잡기–이동–놓기의 pick & place 순서"),
], 3)
doc.add_paragraph("권장 탐구 질문", style="Heading 2")
simple_table(
    ["단계", "교사 질문", "학생이 남길 증거"],
    [
        ("예측", "각도를 10° 바꾸면 끝점은 어느 방향으로 움직일까?", "예상 방향 또는 좌표"),
        ("실험", "한 관절만 바꾸었을 때 무엇이 달라졌나?", "변경값과 결과 좌표"),
        ("경로", "목표점에 도달하는 경로가 하나뿐일까?", "2개 이상의 경로 비교"),
        ("안전", "장애물을 피하려면 어느 중간점이 필요한가?", "경유점과 충돌 여부"),
        ("성찰", "가장 안정적인 경로의 판단 기준은?", "거리·시간·충돌 여유"),
    ],
    [2.5, 8.2, 6.1],
)
callout("주의", "실물 로봇팔은 동작 범위 안에 손을 넣지 않고, 각도·속도 값을 작은 범위에서 단계적으로 바꾸도록 지도합니다.", RED, "FDEEEE")
page_break()

# 11. Mecanum
title("4. 학생 실습 기능", "04 · MECANUM LAB", "메카넘휠의 네 바퀴 회전 방향과 벡터 합을 관찰하고, 주행 경로를 코드로 설계합니다.")
doc.add_paragraph("학습 경로", style="Heading 2")
process_flow(["직진·후진", "좌우 이동", "대각선", "회전", "S자 주행", "블루투스"])
simple_table(
    ["활동", "핵심 과제", "관찰 기준"],
    [
        ("동작 구현", "네 바퀴의 방향을 직접 조합", "각 휠 벡터와 전체 이동 방향"),
        ("S자 코스", "컵 사이 경로와 시간값 설계", "충돌 위치, 회전 시간, 중간 정지"),
        ("블루투스", "명령 문자와 이동 함수 연결", "로봇 기준 앞/옆/회전 방향"),
    ],
    [3.4, 7.0, 6.4],
)
doc.add_paragraph("교사 발문 예시", style="Heading 2")
bullets([
    "오른쪽으로 평행 이동하려면 각 바퀴는 어느 방향으로 회전해야 할까?",
    "두 번째 컵에서 충돌했다면 회전 시간과 직진 시간 중 무엇을 먼저 바꾸겠는가?",
    "여러 값을 동시에 바꾸면 어떤 값이 결과에 영향을 주었는지 확인하기 어려운 이유는 무엇인가?",
])
callout("운영 팁", "S자 코스는 한 번에 전체 경로를 완성하기보다 출발–첫 번째 컵–두 번째 컵 구간으로 나누어 검증하게 합니다.", GREEN, "EAF8F2")
page_break()

# 12. Engineering note
title("4. 학생 실습 기능", "04 · ENGINEERING NOTE", "학생의 문제 해결 과정을 관찰–원인–해결–성찰의 구조로 남기는 핵심 기록 공간입니다.")
screenshot("screen_note.png", "그림 3. 엔지니어링 노트와 AI 노트 코치 비식별 예시")
simple_table(
    ["기록 항목", "좋은 기록의 기준"],
    [
        ("발생한 문제", "보이는 현상을 구체적으로 적음"),
        ("원인 분석", "근거가 있는 원인 후보를 제시함"),
        ("해결 방법", "바꾼 값·순서·조건을 명확히 적음"),
        ("알게 된 점", "예상과 결과의 차이, 다음 실험을 연결함"),
    ],
    [4.0, 12.8],
    8,
)
callout("교사 확인", "‘안 됐다→고쳤다’처럼 결과만 적으면 관찰한 현상, 바꾼 값, 비교 결과를 각각 한 문장씩 보충하게 합니다.", BLUE, LIGHT)
page_break()

# 13. AI coach
title("5. AI 코치 활용법", "05 · QUESTIONING AI", "AI 코치는 학생이 직접 확인할 수 있는 다음 행동을 제안하고, 정답 대신 사고를 이어 주는 질문을 제공합니다.")
doc.add_paragraph("좋은 질문 만들기", style="Heading 2")
simple_table(
    ["덜 좋은 질문", "더 좋은 질문"],
    [
        ("정답 코드 알려줘.", "LED가 켜지지 않는데 핀 11과 코드에서 먼저 비교할 항목은 무엇인가요?"),
        ("왜 안 돼?", "서보가 떨립니다. 전원과 각도값 중 어떤 순서로 확인하면 좋을까요?"),
        ("S자 코스 코드 만들어줘.", "두 번째 컵에서 충돌합니다. 회전 시간만 바꿔 검증하려면 어떻게 비교할까요?"),
        ("노트 써줘.", "원인 분석을 구체화하려면 어떤 관찰값을 더 기록해야 하나요?"),
    ],
    [6.1, 10.7],
)
doc.add_paragraph("학생 질문 공식", style="Heading 2")
process_flow(["목표", "현재 현상", "확인 내용", "바꾼 값", "다음 단계"])
doc.add_paragraph("교사의 개입 기준", style="Heading 2")
bullets([
    "AI 응답을 그대로 복사할 때: 응답 중 실제로 확인할 한 문장을 선택하게 합니다.",
    "같은 질문을 반복할 때: 학생의 관찰 정보가 추가되었는지 확인합니다.",
    "정답만 요구할 때: ‘무엇을 비교해야 하는지’ 또는 ‘어떤 순서로 검사할지’로 질문을 바꾸게 합니다.",
    "AI 응답이 실제 장치와 다를 때: 실물 관찰과 안전 기준을 우선하고 교사가 판단합니다.",
])
callout("원칙", "AI의 제안은 참고 자료입니다. 최종 판단은 실물 관찰, 수업 목표, 안전 기준, 교사의 확인을 따릅니다.", RED, "FDEEEE")
page_break()

# 14. Teacher dashboard
title("6. 교사용 대시보드", "06 · TEACHER PORTAL", "같은 학교 학생의 질문, 노트, 진행, PDF 제출을 한 화면에서 확인하고 필요한 학생에게 먼저 개입합니다.")
screenshot("screen_teacher.png", "그림 4. 교사용 대시보드 비식별 예시")
simple_table(
    ["메뉴", "확인 내용"],
    [
        ("학생 관리", "프로필, 역할, 최근 활동"),
        ("교사 질문", "학생이 교사에게 보낸 질문과 답변 상태"),
        ("AI 질문 기록", "학생 질문과 AI 응답 상세"),
        ("엔지니어링 노트", "활동·문제·원인·해결·성찰"),
        ("학생 피드백", "교사가 남긴 개별 피드백"),
        ("PDF 제출", "학생별 포트폴리오 생성·확인·다운로드"),
        ("학생 개별/전체 분석", "기록 근거를 바탕으로 한 AI 보조 분석"),
    ],
    [4.2, 12.6],
    7.7,
)
callout("우선순위", "답변 대기 질문, 최근 활동이 멈춘 학생, 같은 오류 질문이 반복되는 학생을 먼저 확인합니다.", GREEN, "EAF8F2")
page_break()

# 15. Teacher records
title("6. 교사용 대시보드", "06 · QUESTIONS & RECORDS", "검색과 필터를 이용하면 수업 중 발생한 질문을 학생·학번·모듈·레슨·날짜별로 빠르게 찾을 수 있습니다.")
doc.add_paragraph("교사 질문 처리", style="Heading 2")
bullets([
    "교사 질문 메뉴에서 답변 대기 상태를 우선 확인합니다.",
    "질문 제목과 학생의 실습 맥락을 확인한 뒤 답변합니다.",
    "답변에는 정답 대신 확인 순서, 비교 기준, 안전상 주의를 포함합니다.",
    "처리가 끝난 질문은 상태를 변경해 중복 대응을 줄입니다.",
], numbered=True)
doc.add_paragraph("AI 질문 기록 읽는 법", style="Heading 2")
simple_table(
    ["관찰 신호", "해석 질문", "교사 대응"],
    [
        ("같은 질문 반복", "새 관찰 정보가 추가되었나?", "관찰 항목을 지정해 재실험"),
        ("모듈 전환 잦음", "한 과제를 끝내지 못한 이유는?", "작은 완료 기준 제시"),
        ("정답 요구 중심", "비교·가설 언어가 있는가?", "질문 공식으로 다시 작성"),
        ("구체적 값 포함", "변인 통제가 이루어졌나?", "좋은 탐구 질문으로 강화"),
    ],
    [4.0, 6.2, 6.6],
)
callout("검색 팁", "학생 이름이나 학번으로 범위를 먼저 줄인 뒤 모듈·레슨·날짜를 추가하면 기록을 더 정확히 찾을 수 있습니다.", BLUE, LIGHT)
page_break()

# 16. PDF and feedback
title("7. 피드백과 결과물", "07 · FEEDBACK & PDF", "학생의 엔지니어링 노트와 AI 질문 기록을 과정 중심 포트폴리오로 묶어 수업 결과물로 활용합니다.")
doc.add_paragraph("PDF 제출에서 보는 내용", style="Heading 2")
cards([
    ("엔지니어링 노트", "활동·문제·원인·해결·성찰 기록"),
    ("AI 질문 기록", "학생 질문과 AI 응답, 시간·모듈·레슨"),
    ("선택 구성", "노트만, 질문만, 또는 두 기록 모두"),
], 3)
p("PDF 제출 메뉴는 학생이 별도의 파일을 직접 올린 목록만 의미하지 않습니다. 교사용 대시보드에서 저장된 학습 기록을 학생별 PDF 포트폴리오로 확인하거나 내려받는 기능입니다.")
doc.add_paragraph("피드백 작성 원칙", style="Heading 2")
simple_table(
    ["구성", "예시"],
    [
        ("사실", "회전 시간을 2.0초에서 1.4초로 줄였고 충돌 위치가 달라졌습니다."),
        ("강점", "한 번에 한 값만 바꾸어 결과를 비교한 점이 좋습니다."),
        ("다음 질문", "다음에는 어느 구간의 시간을 0.1초씩 조정해 볼까요?"),
        ("기대 행동", "변경값과 결과를 표로 남기고 최종 선택의 근거를 한 문장으로 적습니다."),
    ],
    [3.2, 13.6],
)
callout("개인정보", "PDF를 외부 제출하거나 공유할 때는 학교의 개인정보 지침에 따라 이름·학번·이메일 포함 여부를 확인합니다.", RED, "FDEEEE")
page_break()

# 17. Individual analysis
title("8. AI 학습 분석", "08 · INDIVIDUAL ANALYSIS", "특정 학생의 AI 질문과 엔지니어링 노트를 근거로 강점, 어려움 신호, 다음 피드백 질문을 정리합니다.")
screenshot("screen_analysis.png", "그림 5. 학생 개별 분석 비식별 예시")
doc.add_paragraph("분석 결과 활용 절차", style="Heading 2")
process_flow(["학생 검색", "분석 실행", "근거 확인", "교사 판단", "다음 질문 선택"])
callout("중요", "분석 문장을 평가 결과로 그대로 사용하지 않습니다. 반드시 원문 기록과 실제 수업 관찰을 확인한 뒤 피드백의 참고 자료로 사용합니다.", RED, "FDEEEE")
page_break()

# 18. Overall analysis
title("8. AI 학습 분석", "08 · WHOLE CLASS", "전체 학생 분석은 학급의 공통 질문 패턴과 어려움 신호를 파악해 다음 차시의 설명, 모둠 구성, 보충 활동을 설계하는 데 사용합니다.")
doc.add_paragraph("확인할 네 가지", style="Heading 2")
cards([
    ("공통 강점", "여러 학생이 잘 수행한 전략"),
    ("공통 어려움", "반복되는 개념·코드·기구 문제"),
    ("기록 습관", "관찰·원인·변경값 기록의 충실도"),
    ("다음 수업", "재설명, 시범, 질문, 보충 과제"),
], 2)
doc.add_paragraph("분석에서 수업 조정으로", style="Heading 2")
simple_table(
    ["분석 신호", "수업 조정 예시"],
    [
        ("핀 번호 불일치 질문이 많음", "회로도와 코드에서 같은 핀을 색으로 표시하는 5분 미니수업"),
        ("여러 값을 동시에 변경", "한 번에 한 변인만 바꾸는 비교 실험표 제공"),
        ("노트에 결과만 기록", "관찰한 현상–바꾼 값–비교 결과 문장 틀 제공"),
        ("S자 코스 두 번째 구간 충돌", "구간별 검증과 중간 정지 전략을 공통 시범"),
        ("AI 정답 요구 질문 증가", "좋은 질문/덜 좋은 질문을 모둠별로 재작성"),
    ],
    [7.2, 9.6],
)
callout("공정성", "질문 수나 노트 길이만으로 학생을 비교하지 않습니다. 활동 기회, 장치 상태, 팀 역할, 교사의 관찰을 함께 고려합니다.", INDIGO, LIGHT)
page_break()

# 19. Admin
title("9. 관리자 기능", "09 · ADMIN SETTINGS", "관리자는 수업 상황에 따라 교사 질문과 AI 코치 사용 여부를 조절하고, 사용할 GPT 모델과 회원 역할을 관리합니다.")
screenshot("screen_admin.png", "그림 6. 관리자 메뉴 비식별 예시")
simple_table(
    ["설정", "기능", "변경 전 확인"],
    [
        ("교사에게 질문하기", "학생 화면의 교사 질문 버튼 켜기/끄기", "수업 중 교사 응답 가능 여부"),
        ("AI 코치 사용하기", "학생 AI 질문·코칭 기능 전체 켜기/끄기", "API 상태, 수업 목표, 예산"),
        ("GPT 모델", "현재 API 키에서 사용할 수 있는 모델 조회·선택", "모델 특성, 비용, 응답 품질"),
        ("회원·역할 관리", "학생·교사·관리자 역할 수정", "대상 계정·학교·권한"),
    ],
    [4.0, 7.2, 5.6],
    7.7,
)
callout("보안", "관리자 권한과 API 키는 공유하지 않습니다. 역할 변경은 대상 계정과 사유를 확인한 뒤 최소 권한 원칙으로 수행합니다.", RED, "FDEEEE")
page_break()

# 20. 50-minute lesson 1
title("10. 실제 수업 운영안", "10 · 50-MINUTE LESSON", "예시: 아두이노 LED 미션. ‘핀 번호 비교’를 중심으로 한 1차시 운영안입니다.")
simple_table(
    ["시간", "수업 단계", "교사 활동", "학생 활동", "관찰 증거"],
    [
        ("0–5분", "도입", "목표·안전·기록 기준 안내", "로그인, 미션 확인", "오늘 조건을 말함"),
        ("5–12분", "예측", "참고 코드와 미션 차이 질문", "달라진 조건 표시", "핀 번호 비교"),
        ("12–22분", "구성", "회로 점검 순서 안내", "LED·저항·핀 연결", "회로 사진/체크"),
        ("22–32분", "코드·실행", "AI 컴파일과 가상 실행 안내", "코드 수정·결과 확인", "오류·변경값"),
        ("32–40분", "문제 해결", "답 대신 확인 질문 제공", "한 값씩 수정·재실험", "비교 결과"),
        ("40–47분", "기록", "노트 문장 틀 제시", "문제·원인·해결 작성", "과정 기록"),
        ("47–50분", "공유", "2개 사례 비교·다음 질문", "성공/실패 근거 공유", "성찰 한 문장"),
    ],
    [1.8, 2.5, 4.2, 4.4, 3.9],
    7.5,
)
doc.add_paragraph("판서 또는 화면 제시 문장", style="Heading 2")
callout("수업 질문", "참고 코드와 오늘 미션에서 달라진 조건은 무엇인가? → 회로와 코드에서 그 조건은 각각 어디에 나타나는가? → 한 가지를 바꾸면 결과가 어떻게 달라지는가?", BLUE, LIGHT)
doc.add_paragraph("최소 성취 기준", style="Heading 2")
bullets([
    "회로 핀과 코드 핀이 일치해야 함을 설명한다.",
    "수정한 값과 실행 결과를 엔지니어링 노트에 남긴다.",
    "AI 답을 복사하지 않고 실제로 확인한 항목을 한 가지 제시한다.",
])
page_break()

# 21. Differentiation
title("10. 실제 수업 운영안", "10 · DIFFERENTIATION", "같은 미션 안에서 지원 수준과 확장 과제를 조절해 학습 속도 차이를 운영합니다.")
doc.add_paragraph("수준별 지원", style="Heading 2")
simple_table(
    ["학생 상태", "교사 지원", "학생 과제"],
    [
        ("시작이 어려움", "회로·코드 점검 체크 3항목 제공", "각 항목을 확인하고 결과 표시"),
        ("오류가 반복됨", "한 번에 한 값만 바꾸도록 제한", "변경 전·후 결과를 표로 비교"),
        ("기본 미션 완료", "새 조건 한 가지 제시", "핀·시간·각도 중 하나를 바꾸고 근거 설명"),
        ("빠른 완성", "최적화 기준 질문", "안정성·속도·정확도 중 기준을 선택해 개선"),
    ],
    [3.3, 6.4, 7.1],
)
doc.add_paragraph("모둠 역할", style="Heading 2")
cards([
    ("구성 담당", "회로·기구 연결과 안전 확인"),
    ("코드 담당", "변경값과 코드 실행 관리"),
    ("기록 담당", "관찰·비교 결과를 노트에 정리"),
    ("검증 담당", "예상과 실제를 비교하고 질문 제안"),
], 2)
callout("운영 팁", "역할은 차시 중 한 번 교대해 모든 학생이 구성·코드·기록·검증을 경험하게 합니다.", GREEN, "EAF8F2")
doc.add_paragraph("수업 종료 전 3분 루틴", style="Heading 2")
process_flow(["실행 중지", "전원 분리", "부품 정리", "노트 저장", "다음 과제 확인"])
page_break()

# 22. Assessment
title("11. 평가 및 관찰", "11 · RUBRIC", "완성 결과보다 문제 해결 과정의 질을 관찰할 수 있도록 네 영역의 간단 루브릭을 사용합니다.")
simple_table(
    ["영역", "3점: 충실", "2점: 부분", "1점: 시작"],
    [
        ("관찰·비교", "조건과 결과를 구체적으로 비교", "일부 차이만 언급", "현상을 모호하게 표현"),
        ("가설·검증", "한 변인을 바꾸고 근거로 판단", "변경은 있으나 기준이 불명확", "여러 값을 동시에 변경"),
        ("과정 기록", "문제·원인·해결·성찰 연결", "일부 항목만 구체적", "결과만 기록"),
        ("AI 활용", "질문을 구체화하고 실제로 확인", "AI 제안 일부를 검증", "정답을 복사·의존"),
        ("협업·안전", "역할을 수행하고 안전 절차 준수", "교사 안내 후 수행", "반복 안내가 필요"),
    ],
    [3.0, 5.1, 4.5, 4.2],
    7.4,
)
doc.add_paragraph("형성평가 관찰 메모", style="Heading 2")
simple_table(
    ["학생/모둠", "발견한 차이", "바꾼 값", "비교 결과", "다음 피드백"],
    [("____________", "________________", "____________", "________________", "________________") for _ in range(4)],
    [2.8, 3.6, 2.7, 3.8, 3.9],
    8,
)
callout("평가 원칙", "AI 분석 결과는 평가 점수를 자동 결정하는 자료가 아닙니다. 실제 활동 관찰과 학생 원문 기록을 중심 증거로 사용합니다.", RED, "FDEEEE")
page_break()

# 23. Safety/privacy
title("12. 안전·개인정보", "12 · SAFETY & PRIVACY", "실물 장치 안전, 계정 권한, 학습 기록 보호를 수업 운영의 기본 조건으로 둡니다.")
doc.add_paragraph("실물 키트 안전", style="Heading 2")
bullets([
    "회로를 바꾸기 전 USB 또는 외부 전원을 먼저 분리합니다.",
    "모터·서보·로봇팔 동작 범위에 손과 얼굴을 넣지 않습니다.",
    "과열, 냄새, 소음, 비정상 진동이 발생하면 즉시 전원을 끕니다.",
    "전압·전류·핀 연결은 부품 사양과 교사 안내를 우선합니다.",
    "실물에서 확인하지 않은 AI 제안을 곧바로 실행하지 않습니다.",
])
doc.add_paragraph("개인정보와 계정", style="Heading 2")
simple_table(
    ["항목", "운영 원칙"],
    [
        ("로그인", "개인 계정을 사용하고 비밀번호를 공유하지 않음"),
        ("학생 정보", "수업 목적에 필요한 최소 정보만 입력·조회"),
        ("PDF·화면 캡처", "외부 공유 전 이름·학번·이메일 노출 여부 확인"),
        ("관리자 권한", "필요한 사람에게만 부여하고 변경 대상 재확인"),
        ("API 키", "브라우저·문서·GitHub에 기록하지 않음"),
        ("AI 입력", "민감한 개인정보나 비공개 자료를 질문에 포함하지 않음"),
    ],
    [4.0, 12.8],
)
callout("수업 안내", "학생에게 ‘AI 입력창도 외부 서비스로 전달되는 공간’임을 설명하고, 개인정보를 입력하지 않도록 반복 안내합니다.", ORANGE, "FFF5E8")
page_break()

# 24. Troubleshooting
title("12. 문제 해결", "12 · TROUBLESHOOTING", "오류가 발생하면 계정–네트워크–저장–기능 설정–실물 장치의 순서로 원인을 분리해 확인합니다.")
simple_table(
    ["증상", "우선 확인", "조치"],
    [
        ("로그인 팝업이 열리지 않음", "인앱 브라우저·팝업 차단", "Chrome/Edge 외부 브라우저로 접속"),
        ("학생 기록이 안 보임", "학교명·계정·저장 여부", "같은 학교명과 로그인 계정 확인 후 재저장"),
        ("AI 응답이 없음", "관리자 AI 상태·네트워크", "AI 코치 켜짐, 새로고침, 잠시 후 재시도"),
        ("교사 질문 버튼 없음", "관리자 기능 설정", "교사에게 질문하기를 켜짐으로 변경"),
        ("GPT 모델 목록이 비어 있음", "API 상태·새로고침", "목록 새로고침 후 사용 가능한 모델 선택"),
        ("시뮬레이션과 실물이 다름", "배선·전원·부품 사양", "실물 관찰을 우선하고 값 하나씩 검증"),
        ("관리/상세 버튼이 안 보임", "화면 폭·브라우저 배율", "배율 100%, 창 최대화, 가로 스크롤 확인"),
        ("저장 후 이전 내용", "캐시·로그인 계정", "강력 새로고침 후 현재 계정 재확인"),
    ],
    [5.4, 5.3, 6.1],
    7.5,
)
doc.add_paragraph("교사 문제 해결 기록", style="Heading 2")
process_flow(["증상 재현", "조건 확인", "조치", "결과 비교", "관리자 전달"])
callout("지원 요청", "문제를 전달할 때 접속 주소, 발생 시각, 역할, 메뉴 이름, 재현 순서, 개인정보를 가린 화면 캡처를 함께 제공합니다.", BLUE, LIGHT)
page_break()

# 25. Checklist
title("부록 A. 수업 운영 체크리스트", "APPENDIX A", "인쇄해 수업 전·중·후 점검표로 사용할 수 있습니다.")
doc.add_paragraph("수업 전", style="Heading 2")
simple_table(
    ["확인", "항목"],
    [
        ("□", "교사 로그인과 교사용 대시보드 접근 확인"),
        ("□", "학생 계정·학교명·역할 확인"),
        ("□", "AI 코치 및 교사 질문 기능 상태 확인"),
        ("□", "키트 수량, 전원, 케이블, 소모품 확인"),
        ("□", "미션·성취 기준·안전 문장 준비"),
    ],
    [1.2, 15.6],
    7.7,
)
doc.add_paragraph("수업 중", style="Heading 2")
simple_table(
    ["확인", "항목"],
    [
        ("□", "참고 조건과 미션 조건의 차이를 말함"),
        ("□", "한 번에 한 값만 바꾸어 비교함"),
        ("□", "AI 응답을 실제 회로·장치와 대조함"),
        ("□", "문제·원인·해결·성찰을 노트에 남김"),
        ("□", "전원 분리와 동작 범위 안전을 지킴"),
    ],
    [1.2, 15.6],
    7.7,
)
doc.add_paragraph("수업 후", style="Heading 2")
simple_table(
    ["확인", "항목"],
    [
        ("□", "답변 대기 질문과 최근 활동 확인"),
        ("□", "노트와 질문 기록을 근거로 피드백 작성"),
        ("□", "전체 분석의 공통 어려움 신호 확인"),
        ("□", "PDF 공유 전 개인정보 노출 확인"),
        ("□", "다음 차시 보충·확장 과제 결정"),
    ],
    [1.2, 15.6],
    7.7,
)

# 26. FAQ + quick ref
title("부록 B. FAQ와 빠른 참조", "APPENDIX B", "수업 중 자주 생기는 질문과 핵심 운영 원칙을 한 쪽에 모았습니다.")
simple_table(
    ["질문", "답"],
    [
        ("AI가 정답을 바로 알려주면 어떻게 하나요?", "학생이 실제로 확인한 항목과 바꾼 값을 말하게 하고, 답을 질문형으로 다시 표현하게 합니다."),
        ("PDF 제출은 무엇을 보는 메뉴인가요?", "저장된 엔지니어링 노트와 AI 질문 기록을 학생별 포트폴리오로 확인·다운로드하는 메뉴입니다."),
        ("분석 결과를 평가에 바로 써도 되나요?", "아닙니다. 원문 기록과 실제 관찰을 확인한 뒤 보조 자료로만 사용합니다."),
        ("관리자 메뉴 비밀번호는 무엇인가요?", "공용 기본 비밀번호가 아니라 지정된 관리자 계정 권한으로 접근합니다."),
        ("사용 가능한 GPT 모델은 어디서 고르나요?", "관리자 메뉴에서 모델 목록을 새로고침한 뒤 드롭다운에서 선택합니다."),
        ("수정 내용이 웹에 바로 안 보이면?", "자동 배포 완료 여부와 브라우저 캐시를 확인하고 강력 새로고침합니다."),
    ],
    [6.1, 10.7],
    7.6,
)
doc.add_paragraph("빠른 참조", style="Heading 2")
cards([
    ("학생에게", "관찰 → 한 값 변경 → 비교 → 기록"),
    ("교사에게", "질문 기록 → 근거 확인 → 다음 질문"),
    ("관리자에게", "기능 상태 → 모델 → 역할 → 보안"),
], 3)
callout("한 문장 요약", "이 자료는 로봇을 완성하는 도구를 넘어, 학생이 자신의 문제 해결 과정을 설명하도록 돕는 AI 기반 수업 플랫폼입니다.", NAVY, LIGHT)
p("접속: https://robot-ai-class.web.app", bold_prefix="접속:", color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
p("운천고등학교 · 유동규 · 제73회 교육자료전 · 2026학년도", color=MUTED, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)


# Document metadata and compatibility
doc.core_properties.title = "생각을 키우는 AI 코치 활용 로봇 수업 자료 교사용 운영 설명서"
doc.core_properties.subject = "제73회 교육자료전 출품용 교사용 설명서"
doc.core_properties.author = "유동규"
doc.core_properties.keywords = "AI 코치, 로봇 수업, 아두이노, 로봇팔, 메카넘, 엔지니어링 노트"

settings = doc.settings._element
update_fields = OxmlElement("w:updateFields")
update_fields.set(qn("w:val"), "true")
settings.append(update_fields)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
